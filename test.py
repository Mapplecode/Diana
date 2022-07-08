from pkgutil import get_data
from bs4 import BeautifulSoup
import requests, urllib.parse
import re
from confi import *
import time
import validators
import docx

Store_Title = search_data

match_keyword_for_neem = "Neem"
fl = "test.docx"
def create_docx(h2list,file_name):
    # print(h2list,get_title)
    doc = docx.Document()
    doc.add_heading('Neem Oil', 0)
    for i in h2list:
        print(i)
        if i[0] != '':
            if i[0] == "h2":

                doc.add_heading(i[1], 3)    
            else:
                doc.add_paragraph(i[1])
                
        doc.add_page_break()
    doc.save(file_name+".docx")
 

def click_url(url,data,input_length,file_name):
    main_list = list()
    h2list = []
    collected = ''
    response = ''
    try:
        response = requests.get(url,headers=headers).text
    except:
        pass
    if validators.url(url) and response != '':
        soup = BeautifulSoup(response, 'html.parser')
        get_len = len(soup.find_all('h2'))
        try:
            if get_len > input_length:
                for tag in soup.find_all():
                    if tag.name == "h2":
                        if data.lower() in tag.text.lower():
                            h2list.append((tag.name,tag.text))
                    elif tag.name == "p":
                        if match_keyword_for_neem.lower() in tag.text.lower():
                            if len(tag.text) > 255:
                                h2list.append((tag.name,tag.text))
                if len(h2list) > 0:
                    main_list = main_list + h2list
                # print("=================",main_list,"=============================")
                create_docx(main_list,file_name)
                return main_list,True
            else:
                return "[INFO] Goto next page..", False
        except Exception as e:
            print(e)
            return "[INFO] Goto next page..", False
    else:
        return "[INFO] Goto next page..", False
def paginate(url,data ,previous_url=None,input_length=300,file_name='file.docx'):
    get_title = ""
    if url == previous_url: return
    headers = {
        "User-Agent":
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/103.0.0.0 Safari/537.36"
    }
    response = requests.get(url, headers=headers).text
    soup = BeautifulSoup(response, 'html.parser')
    links = soup.find_all('div', class_='yuRUbf')
    for link in links:
        title = link.find('h3', text=re.compile(match_keyword_for_neem))
        if title != None:
            get_title = title.text
           
        get_links = link.find('a')
        show_url_list = get_links.get('href')
        print("[INFO] WORKING ON URL : ",show_url_list)
        time.sleep(1)
        datas,status = click_url(show_url_list,data,input_length,file_name)
        # if status:
        #     print(datas,end='\n')
        # else:
        #     print("[INFO] ",show_url_list,"\nURL having 0 H2")
        #     # print(datas,end='\n')
        emp_list.append(show_url_list)

    yield soup
    next_page_node = soup.select_one('a#pnnext')
    if next_page_node is None: return
    next_page_url = urllib.parse.urljoin('https://www.google.com/',next_page_node['href'])
    yield from paginate(next_page_url,data,url,input_length,file_name)

def scrape(data):
    text = input('Enter the text that you want to search.......\n')
    file_name = str(text).replace(' ','_')+'.docx'
    length_needed = 300
    try:
        length_needed = int(input('Enter h2 length needed: \n'))
    except:
        pass
    pages = paginate('https://google.com/search?q='+text,data,input_length=length_needed,file_name=file_name)
    if pages:
        for soup in pages:
            try:
                count_page = int(soup.select_one(".YyVfkd").text)
                if count_page == 1:
                    break
            except Exception as e:
                print(e,soup)
    else:
        print(pages)


def main():
    data = search_data
    for i in data:
        get_keyword = i.get('matchkeyword')
        for show in get_keyword:
            show_all_keyword = show.lower()
            scrape(show_all_keyword)
        
emp_list  = list()
headers = {
    'User-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/103.0.0.0 Safari/537.36'
}
main()
