import requests,os
from bs4 import BeautifulSoup
import aspose.words as aw
from simple_colors import *
import docx
from docx.shared import Inches

url = 'https://www.gardeningknowhow.com/plant-problems/pests/pesticides/neem-oil-uses.htm'
# url ='https://www.medicalnewstoday.com/articles/327179#what-is-neem-oil'

headers = {
    "User-Agent":
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/103.0.0.0 Safari/537.36"
}
response = requests.get(url, headers=headers).text
soup = BeautifulSoup(response, 'html.parser')

h2list = list()
for tag in soup.find_all():
    if tag.name == "h2":
        h2list.append((tag.name,tag.text))
    elif tag.name == "p":
        if len(tag.text) > 255:
            h2list.append((tag.name,tag.text))
    
  
# Create an instance of a word document
doc = docx.Document()
# Add a Title to the document 
doc.add_heading('Neem Oil', 0)

for i in h2list:
    if i[0] == "h2":
        doc.add_heading(i[1], 3)
        # runner.add_break()
    else:
        para = doc.add_paragraph(i[1])
doc.save('gfg.docx')