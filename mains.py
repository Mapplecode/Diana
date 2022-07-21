from bs4 import BeautifulSoup
import requests, urllib.parse
from confi import *
import validators
import docx
	

def click_url(url,data,input_length,filename,doc):
    headers = {
        'User-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/103.0.0.0 Safari/537.36'
    }
    print("Searching Keyword :....",data)
    h2list = []
    response = ''
    try:
        response = requests.get(url,headers=headers).text
    except:
        pass
    if validators.url(url) and response != '':
        soup = BeautifulSoup(response, 'html.parser')
        for tag in soup.find_all('h2'):
            for get_i in data:
                if get_i.lower() in tag.text.lower():
                    GOT_P = False
                    removes = tag.text.find('[')
                    if removes:
                        tag.text[:removes]
                    p_tag = ''
                    try:
                        p_tag = tag.findNext('p')
                        if len(p_tag.text)<=100:
                            p_tag = p_tag.findNext('p')  
                            if p_tag not in [None,'None','']:
                                GOT_P=True
                        if p_tag not in [None,'None','']:
                            GOT_P=True
                        else:
                            p_tag = tag.findNext('p')
          
                        if GOT_P != False:
                            h2list.append(tag.text)
                            h2list.append(p_tag.text[:input_length])
                    except:
                        print("NO result in the 15 links")
        
        if len(h2list) > 0:
            return h2list
    else:
        return "[INFO] Goto next page..", False


def paginate(doc,url,data ,previous_url=None,input_length=300,file_name='file.docx'):
    all_list = []
    if url == previous_url: return
    headers = {
        "User-Agent":
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/103.0.0.0 Safari/537.36"
    }
    response = requests.get(url, headers=headers).text
    soup = BeautifulSoup(response, 'html.parser')
    links = soup.find_all('div', class_='yuRUbf')
    for link in links:
        title = link.find('h3')
        if title != None:
            get_title = title.text
        get_links = link.find('a')
        show_url_list = get_links.get('href')
        if get_links.get('href') not in all_list:
            all_list.append(get_links.get('href'))
            print("[INFO] WORKING ON URL : ",show_url_list)
            store = click_url(show_url_list,data,input_length,file_name,doc)
            emp_list.append(show_url_list)
            main_lists.append(store)

    yield soup
    next_page_node = soup.select_one('a#pnnext')
    if next_page_node is None: return
    next_page_url = urllib.parse.urljoin('https://www.google.com/',next_page_node['href'])
    yield from paginate(doc,next_page_url,data,url,input_length,file_name)

def scrape(get_keyword):
    text = input('Enter the search word that you want \n')
    print(text,"================")
    doc = docx.Document()
    file_name = str(text).replace(' ','_').replace('?','')+'.docx'
    length_needed = 300
    try:
        length_needed = int(input('Enter h2 length needed: \n'))
    except:
        pass
    pages = paginate(doc,'https://google.com/search?q='+text,get_keyword,input_length=length_needed,file_name=file_name)
    if pages:
        for soup in pages:
            try:
                count_page = int(soup.select_one(".YyVfkd").text)
                if count_page == 4:
                    break
            except Exception as e:
                print(e,soup)
    else:
        print(pages)

def main():
    data = search_data
    for i in data:
        get_keyword = i.get('matchkeyword')
        scrape(get_keyword)

emp_list  = list()
main_lists = list()
temp_list = list()
res = []
main()
for i in main_lists:
    if i!=None:
        res.append(i)

doc = docx.Document()
for res_i in res:
    doc.add_heading(res_i[0])
    doc.add_paragraph(res_i[1])    
    doc.add_page_break()
    doc.save('neem.docx')




